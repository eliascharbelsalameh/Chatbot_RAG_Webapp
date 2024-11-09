# document_processing.py

import os
import json
from typing import List
from langchain.docstore.document import Document  # type: ignore
from docx import Document as DocxDocument  # type: ignore
import pdfplumber  # type: ignore
import pandas as pd


def read_files_from_directory(directory_path: str) -> List[Document]:
    """
    Reads and processes files from the specified directory.
    
    Supports PDF, DOCX, XLSX, JSON, and TXT files.
    
    Args:
        directory_path (str): Path to the directory containing files.
        
    Returns:
        List[Document]: A list of processed Document objects.
    """
    supported_extensions = ('.pdf', '.docx', '.xlsx', '.json', '.txt')
    documents = []

    print(f"[Document Processing] Scanning directory: {directory_path}")

    for root, _, files in os.walk(directory_path):
        for file in files:
            if file.lower().endswith(supported_extensions):
                file_path = os.path.join(root, file)
                print(f"[Document Processing] Processing file: {file_path}")
                try:
                    if file.lower().endswith('.pdf'):
                        with pdfplumber.open(file_path) as pdf:
                            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                        print(f"[Document Processing] Successfully processed PDF: {file_path}")

                    elif file.lower().endswith('.docx'):
                        doc = DocxDocument(file_path)
                        text = "\n".join([para.text for para in doc.paragraphs])
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                        print(f"[Document Processing] Successfully processed DOCX: {file_path}")

                    elif file.lower().endswith('.xlsx'):
                        df = pd.read_excel(file_path)
                        text = df.to_string(index=False)
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                        print(f"[Document Processing] Successfully processed XLSX: {file_path}")

                    elif file.lower().endswith('.json'):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        text = json.dumps(data, ensure_ascii=False, indent=4)
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                        print(f"[Document Processing] Successfully processed JSON: {file_path}")

                    elif file.lower().endswith('.txt'):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                        print(f"[Document Processing] Successfully processed TXT: {file_path}")

                except Exception as e:
                    print(f"[Document Processing] Failed to process {file_path}: {e}")

    print(f"[Document Processing] Total documents processed: {len(documents)}")
    return documents
