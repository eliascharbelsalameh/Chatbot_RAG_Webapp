import streamlit as st
import openai
import os, os.path
import pyperclip # type: ignore
from langchain import PromptTemplate # type: ignore
from langchain.chains import RetrievalQA # type: ignore
from langchain.embeddings import HuggingFaceEmbeddings # type: ignore
from langchain.vectorstores import FAISS # type: ignore
from langchain.chat_models import ChatOpenAI # type: ignore
from langchain.text_splitter import CharacterTextSplitter, RecursiveCharacterTextSplitter # type: ignore
from langchain.docstore.document import Document # type: ignore
import pandas as pd
import docx # type: ignore
import tiktoken # type: ignore
import pdfplumber # type: ignore
from gtts import gTTS # type: ignore
import pygame  # for playing the generated audio # type: ignore
from langdetect import detect, LangDetectException # type: ignore

# TODO: relocate vectorstore FAISS to GIN515-Deep Learning
# TODO: add groq-whisper3
# TODO: enhance tables
# TODO: load file online
# TODO: deepgram API for TTS
# TODO: embedding text 3

used_model = "gpt-4-turbo"
chunk_size = 750
chunk_overlap = 75
temperature = 0.1

# Initialize pygame mixer for TTS playback
pygame.mixer.init()

# Define the folder where audio files will be saved
audio_folder = "audio_files"
if not os.path.exists(audio_folder):
    os.makedirs(audio_folder)

def clear_audio_files():
    if pygame.mixer.music.get_busy():
        pygame.mixer.music.stop()  # Stop any audio currently playing
    if os.path.exists(audio_folder):
        for file_name in os.listdir(audio_folder):
            file_path = os.path.join(audio_folder, file_name)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    print(f"Deleted file: {file_path}")
            except Exception as e:
                print(f"Error deleting file {file_path}: {e}")
    else:
        print("Audio folder does not exist.")

clear_audio_files()

def detect_language(text):
    try:
        # Detect the language of the input text
        detected_lang = detect(text)
        return detected_lang
    except LangDetectException:
        # Fallback to English if detection fails
        return 'en'

def play_audio(text, file_name):
    try:
        # Detect the language of the text
        language = detect_language(text)

        audio_file_path = os.path.join(audio_folder, f"{file_name}.mp3")
        if not os.path.exists(audio_file_path):  # Avoid regenerating the audio if it already exists
            tts = gTTS(text, lang=language)  # Use the detected language for TTS
            tts.save(audio_file_path)

        # Load and play the audio using pygame
        pygame.mixer.music.load(audio_file_path)
        pygame.mixer.music.play()

        # Wait until the audio is done playing
        while pygame.mixer.music.get_busy():
            continue

    except Exception as e:
        st.error(f"Error playing audio: {e}")

# Define pages for the Streamlit application
PAGES = ["Chat with Amanda", "Debugging Logs", "User Feedback", "All Chunks"]

# Set up Streamlit sidebar
st.sidebar.title("Menu")
selection = st.sidebar.radio("Go to", PAGES)

# Initialize debugging log, feedback, chunks, and similarity
if "debug_log" not in st.session_state:
    st.session_state["debug_log"] = []
if "feedback" not in st.session_state:
    st.session_state["feedback"] = []  # Store feedback for each response
if "chunks" not in st.session_state:
    st.session_state["chunks"] = []  # Store document chunks
if "last_query" not in st.session_state:
    st.session_state["last_query"] = None  # Store the last query

# Function to log debug messages
def log_debug(message):
    st.session_state["debug_log"].append(message)

# Amanda Page: Main chat interface
if selection == "Chat with Amanda":

    # Set OpenAI API key from environment variable
    openai.api_key = os.getenv("OPENAI_API_KEY")

    # CSS for layout and dark background
    st.markdown(
        """
        <style>
        .reportview-container, .css-1outpf7 {
            background-color: #2c003e;  /* Midnight Purple */
            color: white;
        }
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4, .stMarkdown h5, .stMarkdown h6, .stMarkdown p {
            color: white;
        }
        .stTextInput>div>div {
            background-color: #1e1e1e;
            color: white;
        }
        .stButton>button {
            color: black;
        }
        .stMarkdown {
            color: white;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Title and introductory text
    st.title("🤖 Chat with Amanda (RAG Enhanced)")
    st.write("Amanda is now enhanced with Retrieval-Augmented Generation (RAG) for better, more accurate responses!")

    # Check if the FAISS index files exist and are not corrupted
    def faiss_index_exists():
        index_file = os.path.join(vectorstore_dir, 'index.faiss')
        metadata_file = os.path.join(vectorstore_dir, 'docstore.pkl')
        
        # Check if the necessary files exist
        if not os.path.exists(index_file) or not os.path.exists(metadata_file):
            log_debug("FAISS index or metadata file is missing.")
            return False
        
        try:
            # Attempt to load the files to ensure they are not corrupted
            log_debug("Checking FAISS index and metadata files for corruption...")
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
            FAISS.load_local(vectorstore_dir, embeddings, allow_dangerous_deserialization=True)
            log_debug("FAISS index is valid and not corrupted.")
            return True
        except Exception as e:
            log_debug(f"FAISS index appears corrupted: {e}")
            return False

    # Function to load the vector store, only recreate if it is corrupted or missing
    @st.cache_resource
    def load_vector_store():
        log_debug("Attempting to load FAISS vector store...")
        
        if faiss_index_exists():
            try:
                log_debug("Loading existing FAISS vector store from disk...")
                embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
                vectorstore = FAISS.load_local(vectorstore_dir, embeddings, allow_dangerous_deserialization=True)
                log_debug("Successfully loaded FAISS vector store.")
                return vectorstore
            except Exception as e:
                st.error("Error loading existing FAISS vector store. Recreating...")
                log_debug(f"Error loading FAISS vector store: {e}")

        return recreate_vector_store()

    # Function to recreate the vector store by processing files again
    def recreate_vector_store():
        log_debug("Recreating FAISS vector store from documents...")
        
        documents = read_files_from_directory(input_directory)
        
        log_debug(f"Total documents processed for vector store: {len(documents)}")
        
        if not documents:
            st.error("No documents were loaded. Please check the input directory.")
            log_debug("No documents found. Cannot recreate vector store.")
            return None
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", " "]
        )
        
        texts = text_splitter.split_documents(documents)
        log_debug(f"Text splitter created {len(texts)} chunks.")
        
        if not texts:
            st.error("Text splitting failed. No chunks were created.")
            log_debug("Text splitting failed, no chunks created.")
            return None

        st.session_state["chunks"] = texts
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
        
        try:
            vectorstore = FAISS.from_documents(texts, embeddings)
            vectorstore.save_local(vectorstore_dir)  # Save in the project folder
            log_debug("FAISS vector store created and saved successfully.")
            return vectorstore
        except Exception as e:
            st.error(f"Error creating FAISS vector store: {str(e)}")
            log_debug(f"Error creating FAISS vector store: {str(e)}")
            return None

    # Function to read files from a directory and extract their content
    def read_files_from_directory(directory_path):
        log_debug(f"Reading files from directory: {directory_path}")
        documents = []
        
        try:
            # Check if directory exists
            if not os.path.exists(directory_path):
                st.error(f"The directory {directory_path} does not exist.")
                log_debug(f"The directory {directory_path} does not exist.")
                return documents

            # Traverse the directory
            for dirpath, dirnames, filenames in os.walk(directory_path):
                log_debug(f"Found {len(filenames)} files in {dirpath}")

                for filename in filenames:
                    if filename.startswith('.'):
                        log_debug(f"Skipping hidden file: {filename}")
                        continue

                    file_path = os.path.join(dirpath, filename)
                    log_debug(f"Processing file: {file_path}")

                    try:
                        if filename.lower().endswith('.pdf'):
                            log_debug(f"Extracting content from PDF: {filename}")
                            with pdfplumber.open(file_path) as pdf:
                                for page_number, page in enumerate(pdf.pages, start=1):
                                    text = page.extract_text()
                                    if text:
                                        documents.append(Document(page_content=text, metadata={"source": filename, "page": page_number}))
                                        log_debug(f"Extracted text from PDF: {filename}, Page: {page_number}")
                                    
                                    tables = page.extract_tables()
                                    for table in tables:
                                        table_text = pd.DataFrame(table).to_string()
                                        documents.append(Document(page_content=table_text, metadata={"source": filename, "page": page_number, "type": "table"}))
                                        log_debug(f"Extracted table from PDF: {filename}, Page: {page_number}")

                        elif filename.lower().endswith('.docx'):
                            log_debug(f"Extracting content from Word document: {filename}")
                            doc = docx.Document(file_path)
                            full_text = [para.text for para in doc.paragraphs]
                            documents.append(Document(page_content="\n".join(full_text), metadata={"source": filename}))
                            log_debug(f"Extracted text from Word document: {filename}")

                        elif filename.lower().endswith('.xlsx'):
                            log_debug(f"Extracting content from Excel file: {filename}")
                            df = pd.read_excel(file_path)
                            documents.append(Document(page_content=df.to_string(), metadata={"source": filename}))
                            log_debug(f"Extracted data from Excel file: {filename}")
                        
                        else:
                            log_debug(f"Skipped unsupported file: {filename}")

                    except Exception as e:
                        st.error(f"Error processing file {filename}: {e}")
                        log_debug(f"Error processing file {filename}: {e}")
                        continue

        except Exception as e:
            st.error(f"Error reading files from directory: {e}")
            log_debug(f"Error reading files from directory: {e}")

        log_debug(f"Total documents loaded: {len(documents)}")
        return documents

    # Define the directory where your files are stored
    input_directory = r"C:\Users\elias\OneDrive\Bureau\USEK-Electrical and Electronics Engineer\Semesters\Term-9_Fall-202510\GIN515-Deep Learning-non_repository\Files_dir_RAG"

    # Get the root directory of the project folder dynamically
    project_dir = os.path.dirname(os.path.abspath(__file__))

    # Define a relative path to the 'vectorstore' directory inside your project folder
    vectorstore_dir = os.path.join(project_dir, 'vectorstore/db_faiss')

    def calculate_cost(num_tokens, model=used_model):
        if model == "gpt-4":
            cost_per_1k_tokens = 0.03
        elif model == "gpt-4-turbo":
            cost_per_1k_tokens = 0.012
        elif model == "gpt-3.5-turbo":
            cost_per_1k_tokens = 0.002
        else:
            raise ValueError(f"Unsupported model: {model}")
        cost = (num_tokens / 1000) * cost_per_1k_tokens
        return cost

    def create_rag_chain():
        log_debug("Creating RAG chain...")
        try:
            vectorstore = recreate_vector_store()
            template = """
            Use the following pieces of context to answer the question at the end. 
            If you don't know the answer, just say that you don't know, don't try to make up an answer.
            {context}
            Question: {question}
            Answer: """

            PROMPT = PromptTemplate(template=template, input_variables=["context", "question"])
            qa_chain = RetrievalQA.from_chain_type(
                llm=ChatOpenAI(temperature=temperature, model_name=used_model),
                chain_type="stuff",
                retriever=vectorstore.as_retriever(),
                return_source_documents=True,
                chain_type_kwargs={"prompt": PROMPT}
            )
            log_debug("RAG chain created successfully.")
            return qa_chain
        except Exception as e:
            st.error(f"Error creating RAG chain: {str(e)}")
            log_debug(f"Error creating RAG chain: {str(e)}")
            raise

    qa_chain = create_rag_chain()

    if "messages" not in st.session_state:
        st.session_state["messages"] = [
            {"role": "system", "content": "You are Amanda, a helpful assistant."}
        ]
    
    # Initialize chat-related session state variables
    if "all_chats" not in st.session_state:
        st.session_state["all_chats"] = []  # Holds all chat sessions
    if "active_chat" not in st.session_state:
        st.session_state["active_chat"] = None  # Keeps track of the currently active chat

    # New Chat button
    if st.sidebar.button("New Chat"):
        new_chat_id = len(st.session_state["all_chats"])
        st.session_state["all_chats"].append([])  # Create a new chat session
        st.session_state["active_chat"] = new_chat_id  # Set this as the active chat

    # Display chat selection if multiple chats exist
    if len(st.session_state["all_chats"]) > 1:
        chat_titles = [f"Chat {i + 1}" for i in range(len(st.session_state["all_chats"]))]
        chat_selection = st.sidebar.selectbox("Select Chat", options=chat_titles, index=st.session_state["active_chat"])
        st.session_state["active_chat"] = chat_titles.index(chat_selection)  # Set the selected chat as active

    # Set up the chat messages for the active chat
    if st.session_state["active_chat"] is not None:
        if len(st.session_state["all_chats"]) > st.session_state["active_chat"]:
            active_messages = st.session_state["all_chats"][st.session_state["active_chat"]]
        else:
            active_messages = []
    else:
        active_messages = []

    # Input and response handling for the active chat
    user_input = st.chat_input("Type your message here...")

    if user_input:
        active_messages.append({"role": "user", "content": user_input})
        st.session_state["all_chats"][st.session_state["active_chat"]] = active_messages  # Update the active chat's messages

        with st.spinner("Amanda is thinking..."):
            try:
                # Query the RAG chain to get the response and source documents
                result = qa_chain({"query": user_input})

                # Generate the response with token calculation
                if "result" in result and result["result"]:
                    amanda_message = result["result"].strip()

                    # Extract the most relevant source document metadata
                    source_documents = result.get("source_documents", [])
                    if source_documents:
                        most_relevant_source = source_documents[0]  # Get the most relevant source document
                        metadata = most_relevant_source.metadata
                        filename = metadata.get("source", "Unknown")
                        page = metadata.get("page", "Unknown")
                        chunk_type = metadata.get("type", "text")

                        # Simplified token calculation for cost
                        num_tokens = len(user_input) + len(amanda_message)
                        cost = calculate_cost(num_tokens, model=used_model)

                        # Store the response and its source info along with the cost
                        active_messages.append({
                            "role": "assistant", 
                            "content": amanda_message, 
                            "source": {
                                "filename": filename,
                                "page": page,
                                "type": chunk_type
                            },
                            "tokens": num_tokens,
                            "cost": cost
                        })
                    else:
                        # Handle case where no source documents are found
                        active_messages.append({
                            "role": "assistant", 
                            "content": amanda_message,
                            "tokens": len(user_input) + len(amanda_message),
                            "cost": calculate_cost(len(user_input) + len(amanda_message), model=used_model)
                        })

                # Update the active chat's messages in session state
                st.session_state["all_chats"][st.session_state["active_chat"]] = active_messages

            except Exception as e:
                st.error(f"Error: {e}")

    # Display conversation history for the active chat
    st.markdown("---")
    for idx, message in enumerate(active_messages):
        if message["role"] == "user":
            st.markdown(f"**You:** {message['content']}")
        elif message["role"] == "assistant":
            st.markdown(f"**Amanda 🤖:** {message['content']}")

            # Display the most relevant source information if available
            if "source" in message:
                source_info = message["source"]
                filename = source_info.get("filename", "Unknown")
                page = source_info.get("page", "Unknown")
                chunk_type = source_info.get("type", "text")

                st.markdown(f"""<p style='color: grey;'>Source: File: <i>{filename}</i>, Page: <i>{page}</i></p>""", unsafe_allow_html=True)

            # Display the cost of tokens
            if "cost" in message:
                token_count = message.get("tokens", 0)
                cost = message.get("cost", 0.0)
                st.markdown(f"""<p style='color: grey;'>Tokens used: <i>{token_count}</i>, Cost: <i>${cost:.6f}</i></p>""", unsafe_allow_html=True)

            # Align Play, Like, Dislike, Re-generate, and Copy buttons
            col1, col2, col3, col4, col5 = st.columns([0.1, 0.1, 0.1, 0.1, 0.1])

            with col1:
                play_button = st.button("🔊", key=f"play_{idx}")
                if play_button:
                    play_audio(message['content'], file_name=f"amanda_reply_{idx}")

            with col2:
                like_button = st.button("👍", key=f"like_{idx}")
                if like_button:
                    if len(st.session_state["feedback"]) <= idx:
                        st.session_state["feedback"].append("Liked")
                    else:
                        st.session_state["feedback"][idx] = "Liked"
                    st.success(f"Feedback for message {idx} set to: Liked")

            with col3:
                dislike_button = st.button("👎", key=f"dislike_{idx}")
                if dislike_button:
                    if len(st.session_state["feedback"]) <= idx:
                        st.session_state["feedback"].append("Disliked")
                    else:
                        st.session_state["feedback"][idx] = "Disliked"
                    st.success(f"Feedback for message {idx} set to: Disliked")

            with col4:
                regenerate_button = st.button("🔄", key=f"regenerate_{idx}")
                if regenerate_button:
                    try:
                        result = qa_chain({"query": active_messages[-2]["content"]})
                        amanda_message = result["result"].strip()
                        active_messages[-1]["content"] = amanda_message
                        st.session_state["all_chats"][st.session_state["active_chat"]] = active_messages  # Update after regeneration
                    except Exception as e:
                        st.error(f"Error: {e}")

            with col5:
                copy_button = st.button("📋", key=f"copy_{idx}")
                if copy_button:
                    pyperclip.copy(message['content'])
                    st.success("Copied to clipboard!")

# Debugging Page: Display log messages
elif selection == "Debugging Logs":
    st.title("Debugging Logs")
    st.markdown("### Debugging Information")
    for log_message in st.session_state["debug_log"]:
        st.text(log_message)

# User Feedback Page: Display all feedback collected so far
elif selection == "User Feedback":
    st.title("User Feedback Summary")
    st.markdown("### Feedback Given by Users")
    for i, feedback in enumerate(st.session_state["feedback"]):
        if feedback:
            st.markdown(f"Message {i}: {feedback}")

# Chunks Page: Display chunk settings and all available chunks
elif selection == "All Chunks":
    st.title("Document Chunks")
    st.write(f"**Chunk Size:** {chunk_size}, **Chunk Overlap:** {chunk_overlap}")

    if st.session_state["chunks"]:
        for idx, chunk in enumerate(st.session_state["chunks"]):
            source = chunk.metadata.get("source", "Unknown")
            page = chunk.metadata.get("page", "N/A")
            chunk_type = chunk.metadata.get("type", "text")  # Get the type of chunk (text or table)
            st.markdown(f"**Chunk {idx + 1} from {source}, Page: {page}, Type: {chunk_type}:**")
            st.text(chunk.page_content)
    else:
        st.write("No chunks available. Please load documents.")

# Footer
st.markdown("---")
st.write("© 2024 - Elias-Charbel Salameh and Antonio Haddad")